import tempfile
import unittest
from pathlib import Path

from docx import Document

from tools.sie_autoppt.inputs.source_text import extract_source_text


class SourceTextTests(unittest.TestCase):
    def test_extract_source_text_reads_plain_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "brief.txt"
            path.write_text("Line 1\nLine 2", encoding="utf-8")

            self.assertEqual(extract_source_text(path), "Line 1\nLine 2")

    def test_extract_source_text_strips_html_markup(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "brief.html"
            path.write_text("<html><body><h1>Title</h1><p>Paragraph</p></body></html>", encoding="utf-8")

            self.assertEqual(extract_source_text(path), "Title\nParagraph")

    def test_extract_source_text_reads_docx_paragraphs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "brief.docx"
            doc = Document()
            doc.add_paragraph("Executive Summary")
            doc.add_paragraph("Key milestone and owner")
            doc.save(path)

            self.assertEqual(extract_source_text(path), "Executive Summary\nKey milestone and owner")

